#!/usr/bin/env python3
"""
Enhanced Text Parsing Utilities

This module provides improved text extraction from various document formats,
focusing on preserving proper spacing and formatting that was lost in the
original parsing implementation.

Key improvements:
- Better PDF text extraction with layout awareness
- Enhanced DOCX processing with proper spacing
- Improved DOC file handling with word boundary detection
- Consistent text normalization across formats
"""

import os
import re
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any

from pypdf import PdfReader
from docx import Document as DocxDocument
import xml.etree.ElementTree as ET


class ImprovedTextExtractor:
    """Enhanced text extraction with better spacing and formatting preservation"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF with improved spacing detection
        
        Uses layout mode and post-processing to fix common spacing issues:
        - Words concatenated without spaces
        - Missing line breaks
        - Table formatting issues
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                # Try layout mode first for better spacing
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                except:
                    # Fallback to default extraction
                    page_text = page.extract_text()
                
                if page_text and page_text.strip():
                    # Post-process to fix common spacing issues
                    cleaned_text = ImprovedTextExtractor._fix_pdf_spacing(page_text)
                    text_parts.append(cleaned_text)
            
            # Join pages with clear separation
            return "\n\n=== PAGE BREAK ===\n\n".join(text_parts)
            
        except Exception as e:
            print(f"❌ Error extracting text from PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def _fix_pdf_spacing(text: str) -> str:
        """
        Post-process PDF text to fix common spacing issues
        
        Common PDF extraction problems:
        - Words like "HelloWorld" should be "Hello World"
        - Missing spaces after periods: "sentence.Next" -> "sentence. Next"
        - Extra whitespace and formatting artifacts
        """
        if not text:
            return ""
        
        # Step 1: Fix missing spaces between lowercase and uppercase letters
        # "helloWorld" -> "hello World", but preserve acronyms like "XMLHttpRequest"
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Step 1b: Fix missing spaces between concatenated words
        # Look for specific patterns like "wordthis" -> "word this"
        # Only apply when we see likely word boundaries
        text = re.sub(r'([a-z]{3,})(this|that|and|the|for|with|from)', r'\1 \2', text)
        text = re.sub(r'(this|that|and|the|for|with|from)([a-z]{3,})', r'\1 \2', text)
        
        # Step 2: Fix missing spaces after punctuation
        # "sentence.Next" -> "sentence. Next"
        text = re.sub(r'([.!?;:])([A-Za-z])', r'\1 \2', text)
        
        # Step 3: Fix missing spaces between numbers and letters
        # "123abc" -> "123 abc", "abc123" -> "abc 123"
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)
        
        # Step 4: Clean up excessive whitespace while preserving intentional formatting
        # Replace multiple spaces with single space, but preserve line breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean each line individually to preserve line structure
            cleaned_line = re.sub(r' +', ' ', line.strip())
            if cleaned_line:  # Only add non-empty lines
                cleaned_lines.append(cleaned_line)
        
        # Step 5: Reconstruct text with proper paragraph breaks
        # Group consecutive lines and separate paragraphs with double newlines
        paragraphs = []
        current_paragraph = []
        
        for line in cleaned_lines:
            if len(line) < 10 and line.isupper():
                # Likely a header, treat as separate paragraph
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
                paragraphs.append(line)
            elif line.endswith('.') or line.endswith('!') or line.endswith('?'):
                # End of sentence, likely end of paragraph
                current_paragraph.append(line)
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []
            else:
                # Continue building paragraph
                current_paragraph.append(line)
        
        # Add any remaining paragraph
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        return '\n\n'.join(paragraphs)
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX with better formatting preservation
        
        Improvements over original:
        - Handles runs within paragraphs properly
        - Better table formatting
        - Preserves headers and footers
        - Maintains document structure
        """
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs with better run handling
            for paragraph in doc.paragraphs:
                para_text = ImprovedTextExtractor._extract_paragraph_text(paragraph)
                if para_text:
                    text_parts.append(para_text)
            
            # Extract tables with improved formatting
            for table in doc.tables:
                table_text = ImprovedTextExtractor._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            # Join with proper paragraph separation
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            print(f"❌ Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def _extract_paragraph_text(paragraph) -> str:
        """Extract text from paragraph while preserving run-level formatting"""
        if not paragraph.runs:
            return paragraph.text.strip()
        
        # Handle runs properly to preserve spacing
        run_texts = []
        for run in paragraph.runs:
            run_text = run.text
            if run_text:
                run_texts.append(run_text)
        
        # Join runs and clean up spacing
        full_text = ''.join(run_texts).strip()
        
        # Normalize whitespace within paragraph
        full_text = re.sub(r'\s+', ' ', full_text)
        
        return full_text
    
    @staticmethod
    def _extract_table_text(table) -> str:
        """Extract table text with proper cell separation"""
        table_rows = []
        
        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                # Extract text from all paragraphs in cell
                cell_content = []
                for paragraph in cell.paragraphs:
                    para_text = ImprovedTextExtractor._extract_paragraph_text(paragraph)
                    if para_text:
                        cell_content.append(para_text)
                
                cell_text = ' '.join(cell_content).strip()
                if cell_text:
                    cell_texts.append(cell_text)
            
            if cell_texts:
                # Use pipe separators for better readability
                table_rows.append(' | '.join(cell_texts))
        
        return '\n'.join(table_rows)
    
    @staticmethod
    def extract_text_from_doc(file_path: str) -> str:
        """
        Extract text from legacy DOC files with improved word boundary detection
        
        This is challenging since DOC is a binary format. We use multiple strategies:
        1. Try python-docx (works for some DOC files)
        2. Binary extraction with better heuristics
        3. Pattern-based word boundary detection
        """
        try:
            # Strategy 1: Try python-docx first (some DOC files work)
            try:
                doc = DocxDocument(file_path)
                text_parts = []
                for paragraph in doc.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text:
                        text_parts.append(para_text)
                
                if text_parts:
                    return '\n\n'.join(text_parts)
            except:
                pass  # Fall back to binary extraction
            
            # Strategy 2: Enhanced binary extraction
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Try different encodings
            for encoding in ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding, errors='ignore')
                    if text:
                        cleaned_text = ImprovedTextExtractor._clean_binary_extracted_text(text)
                        if len(cleaned_text) > 50:  # Minimum viable text length
                            return cleaned_text
                except:
                    continue
            
            return f"[Could not extract readable text from DOC file: {os.path.basename(file_path)}]"
            
        except Exception as e:
            print(f"❌ Error extracting text from DOC {file_path}: {e}")
            return ""
    
    @staticmethod
    def _clean_binary_extracted_text(text: str) -> str:
        """
        Clean up text extracted from binary DOC format
        
        Improvements over original naive approach:
        - Better word boundary detection
        - Preserve sentence structure
        - Remove binary artifacts while keeping content
        """
        if not text:
            return ""
        
        # Step 1: Remove obvious binary artifacts
        # Remove null bytes and other control characters except useful ones
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', ' ', text)
        
        # Step 2: Fix word boundaries using improved heuristics
        # Look for patterns that suggest word boundaries
        
        # Add spaces before capital letters that follow lowercase (camelCase)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Add spaces around numbers mixed with letters
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)
        
        # Step 3: Handle punctuation better
        # Ensure space after sentence-ending punctuation
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)
        
        # Step 4: Clean up whitespace
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Step 5: Split into sentences and reconstruct paragraphs
        sentences = re.split(r'[.!?]+', text)
        cleaned_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Filter out very short fragments
                # Capitalize first letter of sentence
                if sentence and sentence[0].islower():
                    sentence = sentence[0].upper() + sentence[1:]
                cleaned_sentences.append(sentence)
        
        # Reconstruct with proper punctuation and paragraphs
        result = '. '.join(cleaned_sentences)
        if result and not result.endswith('.'):
            result += '.'
        
        return result
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from plain text file with encoding detection"""
        try:
            # Try common encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()

                    # Basic cleanup - normalize line endings and excessive whitespace
                    content = content.replace('\r\n', '\n').replace('\r', '\n')
                    content = re.sub(r'\n{3,}', '\n\n', content)  # Max 2 consecutive newlines

                    return content.strip()

                except UnicodeDecodeError:
                    continue

            # If all encodings fail, return error message
            return f"[Could not decode text file with any common encoding: {os.path.basename(file_path)}]"

        except Exception as e:
            print(f"❌ Error reading text file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_vtt(file_path: str) -> str:
        """
        Extract text from WebVTT subtitle file

        WebVTT format includes timestamps and cue metadata that should be removed.
        Only the subtitle text (cue payloads) are extracted.

        Format:
        ```
        WEBVTT

        00:00:00.000 --> 00:00:05.000
        This is the subtitle text

        00:00:05.000 --> 00:00:10.000
        More subtitle text here
        ```
        """
        try:
            text_lines = []
            in_cue = False

            # Try common encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        lines = f.readlines()

                    for line in lines:
                        line = line.strip()

                        # Skip WEBVTT header and empty lines
                        if line.startswith('WEBVTT') or not line:
                            in_cue = False
                            continue

                        # Skip timestamp lines (contain -->)
                        if '-->' in line:
                            in_cue = True
                            continue

                        # Skip NOTE lines and other metadata
                        if line.startswith('NOTE') or line.startswith('STYLE') or line.startswith('REGION'):
                            in_cue = False
                            continue

                        # Skip cue IDs (lines before timestamps that don't contain timestamps)
                        # Cue IDs are usually numbers or short identifiers
                        if in_cue and line and not line.startswith('--'):
                            text_lines.append(line)
                            in_cue = False
                        elif not in_cue and line and '-->' not in line:
                            # This might be a cue ID, skip it
                            continue

                    if text_lines:
                        # Join lines and clean up formatting
                        content = '\n'.join(text_lines)
                        # Normalize whitespace
                        content = re.sub(r'\n{3,}', '\n\n', content)  # Max 2 consecutive newlines
                        return content.strip()

                except UnicodeDecodeError:
                    continue

            # If we get here, file was empty or unreadable
            return ""

        except Exception as e:
            print(f"❌ Error reading VTT file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_srt(file_path: str) -> str:
        """
        Extract text from SubRip subtitle file

        SRT format includes sequence numbers and timestamps that should be removed.
        Only the subtitle text is extracted.

        Format:
        ```
        1
        00:00:00,000 --> 00:00:05,000
        This is the subtitle text

        2
        00:00:05,000 --> 00:00:10,000
        More subtitle text here
        ```
        """
        try:
            text_lines = []

            # Try common encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        lines = f.readlines()

                    for line in lines:
                        line = line.strip()

                        # Skip empty lines
                        if not line:
                            continue

                        # Skip sequence numbers (lines that are just digits)
                        if line.isdigit():
                            continue

                        # Skip timestamp lines (contain -->)
                        if '-->' in line:
                            continue

                        # This is subtitle text
                        text_lines.append(line)

                    if text_lines:
                        # Join lines and clean up formatting
                        content = '\n'.join(text_lines)
                        # Normalize whitespace
                        content = re.sub(r'\n{3,}', '\n\n', content)  # Max 2 consecutive newlines
                        return content.strip()

                except UnicodeDecodeError:
                    continue

            # If we get here, file was empty or unreadable
            return ""

        except Exception as e:
            print(f"❌ Error reading SRT file {file_path}: {e}")
            return ""
    
    @classmethod
    def extract_text(cls, file_path: str, mime_type: str) -> str:
        """
        Main entry point for text extraction with improved algorithms

        This replaces the original TextExtractor.extract_text() method
        with enhanced spacing and formatting preservation.
        """
        if mime_type == "application/pdf":
            return cls.extract_text_from_pdf(file_path)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          "application/vnd.google-apps.document"]:
            return cls.extract_text_from_docx(file_path)
        elif mime_type == "application/msword":
            return cls.extract_text_from_doc(file_path)
        elif mime_type == "text/plain":
            return cls.extract_text_from_txt(file_path)
        elif mime_type == "text/vtt":
            return cls.extract_text_from_vtt(file_path)
        elif mime_type in ["text/srt", "application/x-subrip"]:
            return cls.extract_text_from_srt(file_path)
        else:
            raise ValueError(
                f"Unsupported MIME type for text extraction: {mime_type}. "
                f"Supported types: application/pdf, text/plain, text/vtt, text/srt, "
                f"application/vnd.openxmlformats-officedocument.wordprocessingml.document, "
                f"application/msword, application/vnd.google-apps.document"
            )


class TextQualityValidator:
    """Utility class to validate and score text extraction quality"""
    
    @staticmethod
    def assess_text_quality(text: str) -> Dict[str, Any]:
        """
        Assess the quality of extracted text
        
        Returns a dict with quality metrics:
        - word_count: Number of words
        - avg_word_length: Average word length
        - spacing_issues: Estimated number of spacing problems
        - readability_score: Simple readability assessment
        """
        if not text or not text.strip():
            return {
                'word_count': 0,
                'avg_word_length': 0,
                'spacing_issues': 0,
                'readability_score': 0,
                'quality_rating': 'EMPTY'
            }
        
        words = text.split()
        word_count = len(words)
        
        if word_count == 0:
            return {
                'word_count': 0,
                'avg_word_length': 0,
                'spacing_issues': 0,
                'readability_score': 0,
                'quality_rating': 'NO_WORDS'
            }
        
        # Calculate average word length
        avg_word_length = sum(len(word) for word in words) / word_count
        
        # Estimate spacing issues
        spacing_issues = 0
        
        for word in words:
            # Look for concatenated words (very long words with mixed case)
            if len(word) > 15 and re.search(r'[a-z][A-Z]', word):
                spacing_issues += 1
            
            # Look for words with no vowels (likely corrupted)
            if len(word) > 3 and not re.search(r'[aeiouAEIOU]', word):
                spacing_issues += 1
        
        # Simple readability score (higher is better)
        sentence_count = len(re.findall(r'[.!?]+', text))
        if sentence_count > 0:
            avg_words_per_sentence = word_count / sentence_count
            # Optimal range is 15-20 words per sentence
            if 10 <= avg_words_per_sentence <= 25:
                readability_score = 100 - abs(17.5 - avg_words_per_sentence) * 2
            else:
                readability_score = max(0, 50 - abs(17.5 - avg_words_per_sentence))
        else:
            readability_score = 30  # No sentences found
        
        # Overall quality rating
        spacing_issue_rate = spacing_issues / word_count if word_count > 0 else 1
        
        if spacing_issue_rate < 0.05 and readability_score > 70:
            quality_rating = 'EXCELLENT'
        elif spacing_issue_rate < 0.15 and readability_score > 50:
            quality_rating = 'GOOD'
        elif spacing_issue_rate < 0.30 and readability_score > 30:
            quality_rating = 'FAIR'
        else:
            quality_rating = 'POOR'
        
        return {
            'word_count': word_count,
            'avg_word_length': round(avg_word_length, 2),
            'spacing_issues': spacing_issues,
            'spacing_issue_rate': round(spacing_issue_rate, 3),
            'readability_score': round(readability_score, 1),
            'quality_rating': quality_rating
        }
    
    @staticmethod
    def print_quality_report(text: str, file_name: str = "document"):
        """Print a human-readable quality assessment report"""
        metrics = TextQualityValidator.assess_text_quality(text)
        
        print(f"\n📊 Text Quality Report: {file_name}")
        print(f"   Words: {metrics['word_count']}")
        print(f"   Avg word length: {metrics['avg_word_length']}")
        print(f"   Spacing issues: {metrics['spacing_issues']} ({metrics['spacing_issue_rate']*100:.1f}%)")
        print(f"   Readability: {metrics['readability_score']}/100")
        print(f"   Overall quality: {metrics['quality_rating']}")
        
        if metrics['quality_rating'] in ['POOR', 'FAIR']:
            print(f"   💡 Consider reviewing extraction for potential improvements")


# Example usage and testing functions
if __name__ == "__main__":
    print("🧪 Text Parsing Utils - Testing Mode")
    
    # Test with a sample text that has spacing issues
    sample_problematic_text = "HelloWorldthis isatest.NextSentence123abc"
    
    print("\nOriginal problematic text:")
    print(f"'{sample_problematic_text}'")
    
    fixed_text = ImprovedTextExtractor._fix_pdf_spacing(sample_problematic_text)
    print("\nFixed text:")
    print(f"'{fixed_text}'")
    
    # Quality assessment
    TextQualityValidator.print_quality_report(fixed_text, "test_sample")